from flask import Flask, render_template, request, send_file, jsonify
import os, uuid, sqlite3
from datetime import date
from werkzeug.utils import secure_filename

import img2pdf
import markdown
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
import pandas as pd
import openpyxl
from docx import Document as DocxDocument
from pdf2docx import Converter as PDFConverter
import pytesseract
from PIL import Image
import re
import platform
import shutil
import stripe

# Auto-detect Tesseract on Windows
if platform.system() == 'Windows':
    _found = shutil.which('tesseract')
    if _found:
        pytesseract.pytesseract.tesseract_cmd = _found
    else:
        _win_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            os.path.expanduser(r'~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'),
        ]
        for _p in _win_paths:
            if os.path.exists(_p):
                pytesseract.pytesseract.tesseract_cmd = _p
                break

app = Flask(__name__)


# ── STRIPE CONFIG ─────────────────────────────────────────────────────────────
# Set these in environment variables — never hardcode keys in source
stripe.api_key          = os.environ.get('STRIPE_SECRET_KEY', '')
STRIPE_WEBHOOK_SECRET   = os.environ.get('STRIPE_WEBHOOK_SECRET', '')
STRIPE_PUBLISHABLE_KEY  = os.environ.get('STRIPE_PUBLISHABLE_KEY', '')

# Your Stripe Price IDs — create these in Stripe Dashboard → Products
STRIPE_PRICES = {
    'monthly':  os.environ.get('STRIPE_PRICE_MONTHLY',  ''),
    'yearly':   os.environ.get('STRIPE_PRICE_YEARLY',   ''),
    'lifetime': os.environ.get('STRIPE_PRICE_LIFETIME', ''),
}

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
app.config['UPLOAD_FOLDER']        = os.path.join(BASE_DIR, "uploads")
app.config['OUTPUT_FOLDER']        = os.path.join(BASE_DIR, "outputs")
app.config['MAX_CONTENT_LENGTH']   = 16 * 1024 * 1024
app.config['ALLOWED_EXTENSIONS']   = {
    'docx','xlsx','xls','txt','md','csv',
    'jpg','jpeg','png','gif','bmp','pdf'
}

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

IMAGE_EXTS = {'jpg','jpeg','png','gif','bmp'}

# ── DATABASE ──────────────────────────────────────────────────────────────────
def get_db():
    return sqlite3.connect("database.db")

def init_db():
    db = get_db()
    db.execute("CREATE TABLE IF NOT EXISTS usage (ip TEXT, day TEXT, count INTEGER)")
    db.execute("""CREATE TABLE IF NOT EXISTS conversions
                  (ip TEXT, filename TEXT, from_format TEXT, to_format TEXT,
                   timestamp DATETIME DEFAULT CURRENT_TIMESTAMP)""")
    db.execute("""CREATE TABLE IF NOT EXISTS pro_users
                  (email TEXT PRIMARY KEY,
                   stripe_customer_id TEXT,
                   stripe_subscription_id TEXT,
                   plan TEXT,
                   status TEXT DEFAULT 'active',
                   created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                   expires_at DATETIME)""")
    db.commit()

init_db()

def allowed_file(fn):
    return '.' in fn and fn.rsplit('.',1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def is_pro_user(ip, email=None):
    """Check if IP or email has an active Pro subscription."""
    db = get_db()
    if email:
        row = db.execute(
            "SELECT status FROM pro_users WHERE email=? AND status='active'", (email,)
        ).fetchone()
        if row: return True
    # Also store pro IPs for convenience
    row = db.execute(
        "SELECT status FROM pro_users WHERE stripe_customer_id=? AND status='active'", (ip,)
    ).fetchone()
    return bool(row)

def can_convert(ip, email=None):
    # Pro users have unlimited conversions
    if is_pro_user(ip, email):
        log_conversion(ip, '', '', '')  # still log for stats
        return True
    today = str(date.today())
    db    = get_db()
    row   = db.execute("SELECT count FROM usage WHERE ip=? AND day=?", (ip, today)).fetchone()
    if row and row[0] >= 5:
        return False
    if row:
        db.execute("UPDATE usage SET count=count+1 WHERE ip=? AND day=?", (ip, today))
    else:
        db.execute("INSERT INTO usage VALUES (?,?,1)", (ip, today))
    db.commit()
    return True

def log_conversion(ip, filename, from_f, to_f):
    db = get_db()
    db.execute("INSERT INTO conversions (ip,filename,from_format,to_format) VALUES (?,?,?,?)",
               (ip, filename, from_f, to_f))
    db.commit()

# ── ROUTES ────────────────────────────────────────────────────────────────────
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/upgrade")
def upgrade():
    return render_template("upgrade.html")

@app.route("/stats")
def get_stats():
    ip    = request.remote_addr
    today = str(date.today())
    db    = get_db()
    t = db.execute("SELECT count FROM usage WHERE ip=? AND day=?", (ip, today)).fetchone()
    c = db.execute("SELECT COUNT(*) FROM conversions WHERE ip=?", (ip,)).fetchone()
    return jsonify(today=t[0] if t else 0, total=c[0] if c else 0, limit=5)


# ── STRIPE ROUTES ─────────────────────────────────────────────────────────────

@app.route("/stripe-key")
def stripe_key():
    """Return publishable key to frontend."""
    return jsonify(publishable_key=STRIPE_PUBLISHABLE_KEY)

@app.route("/create-checkout-session", methods=["POST"])
def create_checkout_session():
    """Create a Stripe Checkout session and return the URL."""
    data  = request.get_json()
    plan  = data.get('plan', 'monthly')
    email = data.get('email', '')

    price_id = STRIPE_PRICES.get(plan)
    if not price_id:
        return jsonify(error=f"Invalid plan '{plan}' or price not configured"), 400
    if not stripe.api_key:
        return jsonify(error="Stripe is not configured on the server"), 500

    try:
        # Lifetime = one-time payment, others = subscription
        is_lifetime = plan == 'lifetime'
        mode = 'payment' if is_lifetime else 'subscription'

        params = dict(
            mode=mode,
            line_items=[{'price': price_id, 'quantity': 1}],
            success_url=request.host_url + 'payment-success?session_id={CHECKOUT_SESSION_ID}',
            cancel_url=request.host_url  + 'payment-cancel',
            metadata={'plan': plan},
        )
        if email:
            params['customer_email'] = email

        session = stripe.checkout.Session.create(**params)
        return jsonify(url=session.url, session_id=session.id)

    except stripe.error.StripeError as e:
        return jsonify(error=str(e.user_message)), 400
    except Exception as e:
        return jsonify(error=str(e)), 500


@app.route("/payment-success")
def payment_success():
    """Stripe redirects here after successful payment."""
    session_id = request.args.get('session_id', '')
    session_data = {}

    if session_id and stripe.api_key:
        try:
            session = stripe.checkout.Session.retrieve(
                session_id,
                expand=['customer', 'subscription']
            )
            email    = session.customer_email or (session.customer.email if session.customer else '')
            plan     = session.metadata.get('plan', 'monthly')
            cust_id  = session.customer if isinstance(session.customer, str) else (session.customer.id if session.customer else '')
            sub_id   = ''
            if session.subscription:
                sub_id = session.subscription if isinstance(session.subscription, str) else session.subscription.id

            # Save to DB
            db = get_db()
            db.execute("""
                INSERT INTO pro_users (email, stripe_customer_id, stripe_subscription_id, plan, status)
                VALUES (?, ?, ?, ?, 'active')
                ON CONFLICT(email) DO UPDATE SET
                    stripe_customer_id=excluded.stripe_customer_id,
                    stripe_subscription_id=excluded.stripe_subscription_id,
                    plan=excluded.plan, status='active'
            """, (email, cust_id, sub_id, plan))
            db.commit()

            session_data = {'email': email, 'plan': plan}
        except Exception as e:
            print("payment-success error:", e)

    return render_template("payment_success.html", **session_data)


@app.route("/payment-cancel")
def payment_cancel():
    return render_template("payment_cancel.html")


@app.route("/stripe-webhook", methods=["POST"])
def stripe_webhook():
    """
    Handle Stripe webhook events to keep subscription status in sync.
    Set this URL in Stripe Dashboard → Webhooks:
      https://yourdomain.com/stripe-webhook
    Events to listen for:
      - checkout.session.completed
      - customer.subscription.deleted
      - customer.subscription.updated
      - invoice.payment_failed
    """
    payload = request.get_data()
    sig     = request.headers.get('Stripe-Signature', '')

    if not STRIPE_WEBHOOK_SECRET:
        return jsonify(error="Webhook secret not configured"), 500

    try:
        event = stripe.Webhook.construct_event(payload, sig, STRIPE_WEBHOOK_SECRET)
    except stripe.error.SignatureVerificationError:
        return jsonify(error="Invalid signature"), 400
    except Exception as e:
        return jsonify(error=str(e)), 400

    db = get_db()

    if event['type'] == 'checkout.session.completed':
        session  = event['data']['object']
        email    = session.get('customer_email', '')
        cust_id  = session.get('customer', '')
        sub_id   = session.get('subscription', '')
        plan     = (session.get('metadata') or {}).get('plan', 'monthly')
        if email:
            db.execute("""
                INSERT INTO pro_users (email, stripe_customer_id, stripe_subscription_id, plan, status)
                VALUES (?, ?, ?, ?, 'active')
                ON CONFLICT(email) DO UPDATE SET
                    stripe_customer_id=excluded.stripe_customer_id,
                    stripe_subscription_id=excluded.stripe_subscription_id,
                    plan=excluded.plan, status='active'
            """, (email, cust_id, sub_id, plan))

    elif event['type'] in ('customer.subscription.deleted', 'invoice.payment_failed'):
        obj     = event['data']['object']
        cust_id = obj.get('customer', '')
        if cust_id:
            db.execute("UPDATE pro_users SET status='inactive' WHERE stripe_customer_id=?", (cust_id,))

    elif event['type'] == 'customer.subscription.updated':
        obj     = event['data']['object']
        cust_id = obj.get('customer', '')
        status  = 'active' if obj.get('status') == 'active' else 'inactive'
        if cust_id:
            db.execute("UPDATE pro_users SET status=? WHERE stripe_customer_id=?", (status, cust_id))

    db.commit()
    return jsonify(received=True)

# ── CONVERTERS ────────────────────────────────────────────────────────────────
def _pdf_doc(output_path, margins=50):
    return SimpleDocTemplate(output_path, pagesize=A4,
                              leftMargin=margins, rightMargin=margins,
                              topMargin=margins, bottomMargin=margins)

def docx_to_pdf(inp, out):
    try:
        doc    = DocxDocument(inp)
        styles = getSampleStyleSheet()
        h1  = ParagraphStyle('H1',  parent=styles['Heading1'],  fontSize=18, spaceAfter=10)
        h2  = ParagraphStyle('H2',  parent=styles['Heading2'],  fontSize=14, spaceAfter=8)
        h3  = ParagraphStyle('H3',  parent=styles['Heading3'],  fontSize=12, spaceAfter=6)
        bod = ParagraphStyle('Bod', parent=styles['Normal'],     fontSize=11, leading=16, spaceAfter=5)

        def safe(t):
            return t.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')

        story = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                story.append(Spacer(1, 6))
                continue
            sn = p.style.name if p.style else ''
            if 'Title' in sn or 'Heading 1' in sn:  story.append(Paragraph(safe(txt), h1))
            elif 'Heading 2' in sn:                  story.append(Paragraph(safe(txt), h2))
            elif 'Heading' in sn:                    story.append(Paragraph(safe(txt), h3))
            else:                                    story.append(Paragraph(safe(txt), bod))

        for tbl in doc.tables:
            data = [[c.text.strip() for c in row.cells] for row in tbl.rows]
            if data:
                ncols  = max(len(r) for r in data)
                colw   = (A4[0]-80)/ncols
                t = Table(data, colWidths=[colw]*ncols)
                t.setStyle(TableStyle([
                    ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#6c63ff')),
                    ('TEXTCOLOR',(0,0),(-1,0),colors.white),
                    ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
                    ('FONTSIZE',(0,0),(-1,-1),9),
                    ('GRID',(0,0),(-1,-1),0.4,colors.HexColor('#cccccc')),
                    ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,colors.HexColor('#f5f5ff')]),
                    ('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4),
                ]))
                story += [Spacer(1,8), t, Spacer(1,8)]

        if not story:
            story.append(Paragraph('(empty document)', bod))
        _pdf_doc(out).build(story)
        return True
    except Exception as e:
        print("DOCX→PDF:", e); return False

def images_to_pdf(inp, out):
    try:
        img = Image.open(inp)
        if img.mode in ('RGBA','P','LA'):
            bg = Image.new('RGB', img.size, (255,255,255))
            if img.mode == 'P': img = img.convert('RGBA')
            bg.paste(img, mask=img.split()[-1] if img.mode in ('RGBA','LA') else None)
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        tmp = inp + '_tmp.jpg'
        img.save(tmp, 'JPEG', quality=92)
        with open(out,'wb') as f:
            f.write(img2pdf.convert(tmp))
        os.remove(tmp)
        return True
    except Exception as e:
        print("IMG→PDF:", e)
        try:                        # ReportLab fallback
            img = Image.open(inp).convert('RGB')
            tmp2 = inp+'_rl.jpg'; img.save(tmp2,'JPEG')
            w,h  = A4; iw,ih = img.size
            r    = min(w/iw, h/ih)*0.9; nw,nh = iw*r, ih*r
            c = canvas.Canvas(out, pagesize=A4)
            c.drawImage(tmp2,(w-nw)/2,(h-nh)/2,nw,nh); c.save()
            os.remove(tmp2); return True
        except Exception as e2:
            print("IMG→PDF fallback:", e2); return False

def text_to_pdf(inp, out):
    try:
        with open(inp,'r',encoding='utf-8',errors='replace') as f:
            lines = f.readlines()
        sty = ParagraphStyle('t', fontName='Helvetica', fontSize=10, leading=15, spaceAfter=2)
        story = [Paragraph(l.rstrip().replace('&','&amp;').replace('<','&lt;').replace('>','&gt;') or '&nbsp;', sty)
                 for l in lines]
        _pdf_doc(out).build(story)
        return True
    except Exception as e:
        print("TXT→PDF:", e); return False

def excel_to_pdf(inp, out):
    try:
        wb = openpyxl.load_workbook(inp, data_only=True)
        ws = wb.active
        rows = []
        for row in ws.iter_rows(max_row=200, values_only=True):
            r = [str(c)[:25] if c is not None else '' for c in row]
            if any(r): rows.append(r)
        if not rows:
            rows = [['(empty sheet)']]
        ncols = min(max(len(r) for r in rows), 10)
        rows  = [r[:ncols] for r in rows]
        colw  = (A4[0]-60)/ncols
        tbl   = Table(rows, colWidths=[colw]*ncols, repeatRows=1)
        tbl.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#6c63ff')),
            ('TEXTCOLOR',(0,0),(-1,0),colors.white),
            ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
            ('FONTSIZE',(0,0),(-1,-1),8),
            ('GRID',(0,0),(-1,-1),0.4,colors.HexColor('#dddddd')),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,colors.HexColor('#f7f7ff')]),
            ('TOPPADDING',(0,0),(-1,-1),3),('BOTTOMPADDING',(0,0),(-1,-1),3),
            ('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),4),
        ]))
        _pdf_doc(out, margins=30).build([tbl])
        return True
    except Exception as e:
        print("EXCEL→PDF:", e); return False

def csv_to_pdf(inp, out):
    try:
        df   = pd.read_csv(inp, nrows=200).fillna('')
        hdrs = list(df.columns)
        rows = [hdrs] + [[str(v)[:30] for v in row] for row in df.values.tolist()]
        ncols= min(len(hdrs), 8)
        rows = [r[:ncols] for r in rows]
        colw = (A4[0]-60)/ncols
        tbl  = Table(rows, colWidths=[colw]*ncols, repeatRows=1)
        tbl.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#43d9ad')),
            ('TEXTCOLOR',(0,0),(-1,0),colors.HexColor('#0f1117')),
            ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
            ('FONTSIZE',(0,0),(-1,-1),8),
            ('GRID',(0,0),(-1,-1),0.4,colors.HexColor('#dddddd')),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,colors.HexColor('#f0fdf8')]),
            ('TOPPADDING',(0,0),(-1,-1),3),('BOTTOMPADDING',(0,0),(-1,-1),3),
            ('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),4),
        ]))
        _pdf_doc(out, margins=30).build([tbl])
        return True
    except Exception as e:
        print("CSV→PDF:", e); return False

def markdown_to_pdf(inp, out):
    try:
        with open(inp,'r',encoding='utf-8',errors='replace') as f:
            html = markdown.markdown(f.read())
        text = re.sub(r'<br\s*/?>', '\n', html)
        text = re.sub(r'</p>', '\n', text)
        text = re.sub(r'<[^>]+>', '', text)
        sty  = ParagraphStyle('m', fontName='Helvetica', fontSize=11, leading=16, spaceAfter=4)
        def safe(t):
            return t.replace('&amp;','&').replace('&lt;','<').replace('&gt;','>')\
                    .replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
        story = [Paragraph(safe(l.strip()) or '&nbsp;', sty) for l in text.splitlines()]
        _pdf_doc(out).build(story)
        return True
    except Exception as e:
        print("MD→PDF:", e); return False

def pdf_to_docx(inp, out):
    try:
        cv = PDFConverter(inp)
        cv.convert(out)
        cv.close()
        return True
    except Exception as e:
        print("PDF→DOCX:", e); return False

def image_to_text(inp, out):
    try:
        img = Image.open(inp).convert('RGB')
        # Pre-process: slightly sharpen for better OCR accuracy
        from PIL import ImageFilter
        img = img.filter(ImageFilter.SHARPEN)
        text = pytesseract.image_to_string(img, config='--psm 6')
        if not text.strip():
            text = "(No text detected in image)"
        with open(out, 'w', encoding='utf-8') as f:
            f.write(text)
        return True
    except pytesseract.TesseractNotFoundError:
        print("OCR: Tesseract not found")
        # Write a helpful error file instead of failing silently
        with open(out, 'w', encoding='utf-8') as f:
            f.write("ERROR: Tesseract OCR is not installed or not found.\n\n"
                    "To fix this on Windows:\n"
                    "1. Download from: https://github.com/UB-Mannheim/tesseract/wiki\n"
                    "2. Install it (default path: C:\\Program Files\\Tesseract-OCR)\n"
                    "3. Restart your Flask server\n")
        return True  # Return True so user gets the error file, not a 500
    except Exception as e:
        print("OCR:", e); return False

# ── MAIN CONVERT ENDPOINT ─────────────────────────────────────────────────────
CONVERTERS = {
    'docx': docx_to_pdf,
    'jpg': images_to_pdf, 'jpeg': images_to_pdf,
    'png': images_to_pdf, 'gif':  images_to_pdf, 'bmp': images_to_pdf,
    'txt': text_to_pdf,
    'xlsx': excel_to_pdf, 'xls': excel_to_pdf,
    'csv':  csv_to_pdf,
    'md':   markdown_to_pdf,
    'pdf':  pdf_to_docx,
}

@app.route("/convert", methods=["POST"])
def convert_file():
    if 'file' not in request.files:
        return jsonify(error="No file uploaded"), 400
    file = request.files['file']
    if not file.filename:
        return jsonify(error="No file selected"), 400
    if not allowed_file(file.filename):
        return jsonify(error="File type not allowed"), 400

    ip = request.remote_addr
    if not can_convert(ip):
        return jsonify(error="Daily limit reached. You have used 5 free conversions today. Please try again tomorrow."), 429

    uid        = str(uuid.uuid4())[:8]
    input_name = secure_filename(file.filename)
    ext        = input_name.rsplit('.',1)[1].lower()
    base       = input_name.rsplit('.',1)[0]
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uid}_{input_name}")
    file.save(input_path)

    conversion_type = request.form.get('conversion_type', '')

    if ext == 'pdf':
        out_ext = 'docx'
    elif ext in IMAGE_EXTS and conversion_type == 'img_to_txt':
        out_ext = 'txt'
    else:
        out_ext = 'pdf'

    output_name = f"{base}_{uid}.{out_ext}"
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_name)

    if ext in IMAGE_EXTS and conversion_type == 'img_to_txt':
        fn = image_to_text
    else:
        fn = CONVERTERS.get(ext)

    try:
        if fn is None:
            os.remove(input_path)
            return jsonify(error="Unsupported format"), 400

        ok = fn(input_path, output_path)

        if os.path.exists(input_path):
            os.remove(input_path)

        if not ok or not os.path.exists(output_path):
            return jsonify(error="Conversion failed — please check your file and try again."), 500

        log_conversion(ip, input_name, ext, out_ext)

        return send_file(output_path, as_attachment=True, download_name=output_name)

    except Exception as e:
        if os.path.exists(input_path):
            os.remove(input_path)
        print("UNHANDLED:", e)
        return jsonify(error=f"Unexpected error: {str(e)}"), 500

if __name__ == "__main__":
    app.run(debug=True)
